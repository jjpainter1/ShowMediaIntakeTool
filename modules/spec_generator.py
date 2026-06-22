"""Generate a filled-in delivery specification Word document from show config."""

from __future__ import annotations

import re
from copy import deepcopy
from pathlib import Path

import docx
from docx.table import Table, _Row

from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

from modules.config import ShowConfig, is_flat_intake, is_per_screen_output
from modules.filename_parser import (
    DEFAULT_TOKENS,
    build_example_filename,
    build_filename_pattern,
    example_value_for_token,
)

_TEMPLATES_DIR = Path(__file__).parent.parent / "templates"
_TEMPLATE_PATH = _TEMPLATES_DIR / "spec_template.docx"

_COLOR_SPACE_NAMES: dict[str, str] = {
    "bt709": "Rec.709",
    "bt2020": "Rec.2020",
    "smpte170m": "Rec.601",
    "bt470bg": "Rec.601",
}

_CODEC_DISPLAY_NAMES: dict[str, str] = {
    "prores_422_proxy": "ProRes 422 Proxy",
    "prores_422_lt": "ProRes 422 LT",
    "prores_422": "ProRes 422",
    "prores_422_hq": "ProRes 422 HQ",
    "prores_4444": "ProRes 4444",
    "prores_4444_xq": "ProRes 4444 XQ",
    "notchlc": "NotchLC",
    "h264": "H.264",
    "h265": "H.265 / HEVC",
    "dnxhd": "DNxHD / DNxHR",
    "mpeg4": "MPEG-4",
    "wmv3": "WMV3",
}

_COLOR_RANGE_TEXT: dict[str, str] = {
    "tv": "video/legal range (luma 16-235). Not full range",
    "pc": "full range",
}

_LEGACY_PATTERN = "SCR##_ContentSlug_v##_YYYYMMDD.ext"
_LEGACY_FILENAME_TOKENS = tuple(DEFAULT_TOKENS)
_TEMPLATE_AUDIO_TEXT = (
    "48kHz, 24-bit, stereo (2.0). Audio extracted at intake - embed in videos, "
    "don't deliver separately"
)
_TV_RANGE_SUFFIX = "video/legal range (luma 16-235). Not full range"

_TOKEN_FORMAT: dict[str, str] = {
    "show_token": "Show identifier token configured for this delivery",
    "initials": "2-3 letter artist initials (letters only)",
    "screen": "SCR + 2-digit screen number, audience POV L→R",
    "content": "PascalCase, no spaces, special characters limited to “-“ and “_”",
    "version": "Numeric only. Increments only. Never use \"final\" in any form",
    "date": "Render/export date",
}

_INTAKE_NOTE_ROUTED = (
    "Delivery organization: deliver content in separate folders by screen as shown above. "
    "Filenames must include the correct screen prefix so files route to the intended screen."
)
_INTAKE_NOTE_FLAT = (
    "Delivery organization: deliver all files in a single folder. Filenames are preserved as "
    "delivered; the receiving operator assigns content to screens during intake."
)

DEFAULT_SCREEN_NOTES = (
    "See attached screen diagram for physical placement and template files for each "
    "screen's exact pixel dimensions"
)

_NOTES_COLUMN_PLACEHOLDER = "[Optional Notes]"


def generate_spec(show_root: Path, config: ShowConfig) -> Path:
    """Fill the spec template from *config* and save to the show root."""
    if not _TEMPLATE_PATH.exists():
        raise FileNotFoundError(f"Spec template not found: {_TEMPLATE_PATH}")

    doc = docx.Document(str(_TEMPLATE_PATH))

    _apply_header_replacements(doc, config)
    _apply_literal_replacements(doc, config)
    _rebuild_screens_table(doc.tables[0], config)
    _fill_optional_screen_notes_paragraph(doc, config)
    _rebuild_folder_structure(doc, config)
    _update_video_specs_table(doc.tables[1], config)
    _rebuild_filename_convention_table(doc.tables[2], config)
    _apply_supplemental_sections(doc, config)

    output_path = show_root / f"{config.show_name}_DeliverySpec.docx"
    doc.save(str(output_path))
    return output_path


def _apply_header_replacements(doc: docx.Document, config: ShowConfig) -> None:
    """Replace project metadata placeholders across body, tables, headers, and footers."""
    company_name = config.operator.company_name or "[Company Name]"
    for para in _iter_document_paragraphs(doc):
        if "[Company Name]" in para.text:
            _replace_company_name(para, company_name)

    replacements = {
        "[Project Name]": config.show_name,
        "[##]": str(len(config.screens)),
        "[Operator Name]": config.operator.name,
        "[email@prestigeav.com]": config.operator.email,
    }
    for para in _iter_document_paragraphs(doc):
        for old, new in replacements.items():
            if old in para.text:
                _replace_in_paragraph(para, {old: new})

    for para in _iter_document_paragraphs(doc):
        if "[YYYY-MM-DD]" in para.text:
            _replace_in_paragraph(para, {"[YYYY-MM-DD]": config.show_date}, count=1)


def _replace_company_name(para, name: str) -> None:
    """Replace company name and apply default (non-placeholder) body color."""
    _replace_in_paragraph(para, {"[Company Name]": name})
    for run in para.runs:
        if name in run.text:
            _clear_run_color(run)


def _apply_literal_replacements(doc: docx.Document, config: ShowConfig) -> None:
    """Replace fixed template example lines with config-driven text."""
    pattern = build_filename_pattern(config)
    example = build_example_filename(config)
    date_compact = _show_date_yyyymmdd(config)

    for para in doc.paragraphs:
        text = para.text
        if _LEGACY_PATTERN in text or _is_filename_pattern_line(text):
            _set_paragraph_text_preserve_style(para, pattern)
        elif text.startswith("Full example:"):
            _set_full_example_line(para, example)
        elif "[Project Name]" in text and "YYYYMMDD" in text:
            _replace_in_paragraph(para, {"[Project Name]": config.show_name})
            _replace_in_paragraph(para, {"YYYYMMDD": date_compact})


def _apply_supplemental_sections(doc: docx.Document, config: ShowConfig) -> None:
    """Add intake mode note and optional vendor notes."""
    _insert_intake_mode_note(doc, config)
    _apply_flat_intake_key_rules(doc, config)
    _insert_vendor_notes_section(doc, config)


def _apply_flat_intake_key_rules(doc: docx.Document, config: ShowConfig) -> None:
    """Replace routed folder rule when flat intake omits the folder hierarchy."""
    if not is_flat_intake(config):
        return
    for para in doc.paragraphs:
        if "Match the folder structure" in para.text:
            _set_paragraph_text_preserve_style(
                para,
                "Deliver all files in a single folder. Filenames are preserved as delivered.",
            )
            return


def _insert_intake_mode_note(doc: docx.Document, config: ShowConfig) -> None:
    if is_flat_intake(config):
        return

    note = _INTAKE_NOTE_ROUTED
    placeholder = "[Intake Mode Note]"
    for para in doc.paragraphs:
        if placeholder in para.text:
            _replace_in_paragraph(para, {placeholder: note})
            return

    anchor_idx = _paragraph_index(doc, lambda text: text.startswith("└──"))
    if anchor_idx is None:
        anchor_idx = _paragraph_index(doc, lambda text: "BEFORE FULL DELIVERY" in text)
        if anchor_idx is not None and anchor_idx > 0:
            anchor_idx -= 1
    if anchor_idx is None:
        return
    anchor = doc.paragraphs[anchor_idx]
    new_para = _insert_paragraph_after(anchor)
    _write_paragraph_from_template(new_para, note, anchor)


def _insert_vendor_notes_section(doc: docx.Document, config: ShowConfig) -> None:
    notes = (config.delivery.vendor_notes or "").strip()
    if not notes:
        return

    placeholder = "[Vendor Notes]"
    for para in doc.paragraphs:
        if placeholder in para.text:
            _replace_in_paragraph(para, {placeholder: notes})
            return

    key_rules_idx = _paragraph_index(doc, lambda text: "KEY RULES" in text)
    if key_rules_idx is None:
        return
    key_rules_para = doc.paragraphs[key_rules_idx]
    header = _insert_paragraph_before(key_rules_para)
    _write_paragraph_from_template(header, "  VENDOR NOTES", key_rules_para)
    body = _insert_paragraph_before(key_rules_para)
    style_source = doc.paragraphs[4] if len(doc.paragraphs) > 4 else key_rules_para
    _write_paragraph_from_template(body, notes, style_source)


def _paragraph_index(doc: docx.Document, predicate) -> int | None:
    for index, para in enumerate(doc.paragraphs):
        if predicate(para.text):
            return index
    return None


def _delete_paragraph(para: Paragraph) -> None:
    """Remove a paragraph element from the document body."""
    element = para._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def _insert_paragraph_after(paragraph) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    return Paragraph(new_p, paragraph._parent)


def _insert_paragraph_before(paragraph) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addprevious(new_p)
    return Paragraph(new_p, paragraph._parent)


def _write_paragraph_from_template(target, text: str, style_source) -> None:
    if style_source.runs:
        run = target.add_run(text)
        _copy_run_format(style_source.runs[0], run)
    else:
        target.add_run(text)


def _update_video_specs_table(table: Table, config: ShowConfig) -> None:
    """Fill codec, framerate, color space, and audio rows via targeted placeholders."""
    for row in table.rows:
        if len(row.cells) < 2:
            continue
        label = row.cells[0].text.strip()
        para = row.cells[1].paragraphs[0]
        if label == "Frame Rate":
            _replace_framerate_cell(para, config)
        elif label == "Codec":
            _replace_in_paragraph(para, {"[Codec]": _format_codec_list(config)})
        elif label == "Color Space":
            _replace_color_space_cell(para, config)
        elif label == "Audio":
            _replace_audio_cell(para, config)


def _rebuild_screens_table(table: Table, config: ShowConfig) -> None:
    """Replace example screen rows; always retain the SCRall row from the template."""
    if len(table.rows) < 2:
        return

    scrall_template_row: _Row | None = None
    for row in table.rows[1:]:
        if row.cells[0].text.strip() == "SCRall":
            scrall_template_row = row
            break

    template_row = table.rows[1]
    while len(table.rows) > 1:
        table._tbl.remove(table.rows[-1]._tr)

    for screen in config.screens:
        row = _append_cloned_row(table, template_row)
        _fill_configured_screen_row(row, screen, template_row)

    scrall_row = (
        _append_cloned_row(table, scrall_template_row)
        if scrall_template_row is not None
        else _append_cloned_row(table, template_row)
    )
    _fill_scrall_row(scrall_row, scrall_template_row or template_row)


def _fill_configured_screen_row(row: _Row, screen, template_row: _Row) -> None:
    _write_cell_from_template(row.cells[0], screen.id, template_row.cells[0])
    for placeholder in ("[e.g. House Left]", "[e.g. Center]", "[e.g. House Right]"):
        if placeholder in row.cells[1].text:
            _replace_in_cell(row.cells[1], {placeholder: screen.name or ""})
            break
    else:
        _write_cell_from_template(row.cells[1], screen.name or "", template_row.cells[1])
    _replace_in_cell(row.cells[2], {"[#### x ####]": screen.resolution or "N/A"})
    _set_notes_column_placeholder(row.cells[3], template_row.cells[3])


def _fill_scrall_row(row: _Row, template_row: _Row) -> None:
    _write_cell_from_template(row.cells[0], "SCRall", template_row.cells[0])
    for placeholder in ("[e.g. All Screens]", "[e.g. House Left]"):
        if placeholder in row.cells[1].text:
            _replace_in_cell(row.cells[1], {placeholder: "All Screens"})
            break
    else:
        _write_cell_from_template(row.cells[1], "All Screens", template_row.cells[1])
    _replace_in_cell(row.cells[2], {"[#### x ####]": "N/A"})
    _set_notes_column_placeholder(row.cells[3], template_row.cells[3])


def _set_notes_column_placeholder(cell, template_cell) -> None:
    """Leave the Notes column for manual operator entry after generation."""
    _write_cell_from_template(cell, _NOTES_COLUMN_PLACEHOLDER, template_cell)


def _fill_optional_screen_notes_paragraph(doc: docx.Document, config: ShowConfig) -> None:
    """Fill the single optional-notes block below the screen configuration table."""
    configured = (config.delivery.optional_screen_notes or "").strip()
    text = configured or DEFAULT_SCREEN_NOTES
    for para in doc.paragraphs:
        if "[Optional Notes]" in para.text:
            _replace_in_paragraph(para, {"[Optional Notes]": text})
            return


def _rebuild_folder_structure(doc: docx.Document, config: ShowConfig) -> None:
    """Rebuild folder tree for routed intake; compact flat layout in folder section."""
    if is_flat_intake(config):
        _apply_flat_folder_structure(doc)
        return

    lines = _folder_structure_lines(config)
    folder_indices = [
        index
        for index, para in enumerate(doc.paragraphs)
        if para.text.startswith("├──") or para.text.startswith("└──")
    ]
    if not folder_indices:
        return

    first_idx = folder_indices[0]
    template_para = doc.paragraphs[first_idx]
    for offset, line in enumerate(lines):
        idx = first_idx + offset
        if idx < len(doc.paragraphs):
            _set_paragraph_text_preserve_style(doc.paragraphs[idx], line)
        else:
            new_para = _insert_paragraph_after(doc.paragraphs[idx - 1])
            _write_paragraph_from_template(new_para, line, template_para)

    extra_start = first_idx + len(lines)
    for idx in folder_indices[len(lines) :]:
        if idx < len(doc.paragraphs):
            _set_paragraph_text_preserve_style(doc.paragraphs[idx], "")


def _apply_flat_folder_structure(doc: docx.Document) -> None:
    """Keep the folder section header; place flat delivery text in the grey content area."""
    tree_indices = [
        index
        for index, para in enumerate(doc.paragraphs)
        if para.text.startswith("├──") or para.text.startswith("└──")
    ]
    if not tree_indices:
        return

    root_idx = tree_indices[0] - 1
    if root_idx >= 0:
        _set_paragraph_text_preserve_style(doc.paragraphs[root_idx], _INTAKE_NOTE_FLAT)

    for index in reversed(tree_indices):
        _delete_paragraph(doc.paragraphs[index])


def _folder_structure_lines(config: ShowConfig) -> list[str]:
    example = build_example_filename(config)
    reference_screen = config.screens[0].id if config.screens else "SCR01"
    entries: list[tuple[str, str, str]] = []

    for index, screen in enumerate(config.screens, start=1):
        folder = f"{index:02d}_{screen.id}"
        file_example = example.replace(reference_screen, screen.id)
        entries.append((folder, screen.id, file_example))

    scrall_index = len(config.screens) + 1
    scrall_example = example.replace(reference_screen, "SCRall")
    entries.append((f"{scrall_index:02d}_SCRall", "SCRall", scrall_example))

    aud_index = len(config.screens) + 2
    aud_example = _example_aud_filename(example, reference_screen)
    entries.append((f"{aud_index:02d}_AUD", "AUD", aud_example))

    lines: list[str] = []
    for idx, (folder, _prefix, file_example) in enumerate(entries):
        branch = "└──" if idx == len(entries) - 1 else "├──"
        lines.append(f"{branch} {folder}/ {file_example}")
    return lines


def _example_aud_filename(example: str, reference_screen: str) -> str:
    if reference_screen and example.startswith(f"{reference_screen}_"):
        return f"AUD_{example[len(reference_screen) + 1:]}"
    slug = "AudioExample"
    if "_v" in example:
        suffix = example.split("_v", 1)[1]
        return f"AUD_{slug}_v{suffix}"
    return "AUD_AudioExample_v01_20260425.mov"


def _rebuild_filename_convention_table(table: Table, config: ShowConfig) -> None:
    """Rebuild filename token rows using raw token names and saved token order."""
    if len(table.rows) < 2:
        return
    template_row = table.rows[1]
    convention = config.filename_convention
    tokens = convention.tokens if convention.enabled else list(_LEGACY_FILENAME_TOKENS)

    while len(table.rows) > 1:
        table._tbl.remove(table.rows[-1]._tr)

    for token in tokens:
        row = _append_cloned_row(table, template_row)
        fmt = _TOKEN_FORMAT.get(token, "")
        if token == "content" and convention.allow_loop_suffix:
            fmt += ". Append -LOOP to the slug for seamless loop content"
        _write_cell_from_template(row.cells[0], token, template_row.cells[0])
        _write_cell_from_template(row.cells[1], fmt, template_row.cells[1])
        example = example_value_for_token(token, config)
        _write_cell_from_template(row.cells[2], example, template_row.cells[2])


def _format_codec_list(config: ShowConfig) -> str:
    preferred = [_codec_display_name(codec) for codec in config.preferred_codecs]
    return ", ".join(preferred)


def _replace_framerate_cell(para, config: ShowConfig) -> None:
    if is_per_screen_output(config):
        _replace_in_paragraph(
            para,
            {
                "[e.g. 30 / 60 fps]": (
                    "Varies by screen — see Notes in the Screen Configuration table"
                ),
            },
        )
        return
    framerate = config.expected_specs.framerate
    if framerate is None:
        _replace_in_paragraph(para, {"[e.g. 30 / 60 fps]": "As agreed with operator"})
        return
    _replace_in_paragraph(para, {"[e.g. 30 / 60 fps]": f"{framerate:g}"})


def _replace_color_space_cell(para, config: ShowConfig) -> None:
    if is_per_screen_output(config):
        _set_paragraph_text_preserve_style(
            para,
            "Varies by screen — see Notes in the Screen Configuration table",
        )
        return
    color_space = config.expected_specs.color_space
    color_range = config.expected_specs.color_range
    label = _color_space_label(color_space) if color_space else "As agreed with operator"
    _replace_in_paragraph(para, {"[Color Space]": label})
    if color_range == "pc":
        _replace_in_paragraph(para, {_TV_RANGE_SUFFIX: _COLOR_RANGE_TEXT["pc"]})


def _replace_audio_cell(para, config: ShowConfig) -> None:
    sample_rate = config.expected_specs.audio_sample_rate
    channels = config.expected_specs.audio_channels
    if sample_rate is None and channels is None:
        return
    rate_khz = f"{sample_rate // 1000}kHz" if sample_rate else "As agreed"
    channel_text = _audio_channels_label(channels) if channels else "As agreed"
    _set_paragraph_text_preserve_style(
        para,
        f"{rate_khz}, 24-bit, {channel_text}. Audio extracted at intake - embed in "
        "videos, don't deliver separately",
    )


def _codec_display_name(codec_id: str) -> str:
    return _CODEC_DISPLAY_NAMES.get(codec_id, codec_id.replace("_", " ").title())


def _color_space_label(value: str | None) -> str:
    if not value:
        return "As agreed with operator"
    return _COLOR_SPACE_NAMES.get(value, value)


def _color_range_label(value: str) -> str:
    return _COLOR_RANGE_TEXT.get(value, value)


def _audio_channels_label(channels: int) -> str:
    if channels == 1:
        return "mono (1.0)"
    if channels == 2:
        return "stereo (2.0)"
    return f"{channels} channels"


def _show_date_yyyymmdd(config: ShowConfig) -> str:
    return config.show_date.replace("-", "")


def _is_filename_pattern_line(text: str) -> bool:
    if not text.endswith(".ext"):
        return False
    if re.fullmatch(r"(\{[a-z_]+\}_?)+\.ext", text):
        return True
    return bool(re.fullmatch(r"[a-z_]+\.ext", text))


def _set_full_example_line(para, example: str) -> None:
    if len(para.runs) >= 2:
        para.runs[-1].text = example
        return
    _replace_in_paragraph(para, {para.text: f"Full example: {example}"})


def _set_paragraph_text_preserve_style(para, text: str) -> None:
    if not para.runs:
        para.add_run(text)
        return
    para.runs[0].text = text
    for run in para.runs[1:]:
        run.text = ""


def _write_cell_from_template(cell, text: str, template_cell) -> None:
    """Write cell text using the first run style from a template cell."""
    template_para = template_cell.paragraphs[0]
    para = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    template_run = template_para.runs[0] if template_para.runs else None
    if para.runs:
        target_run = para.runs[0]
    else:
        target_run = para.add_run("")
    if template_run is not None:
        _copy_run_format(template_run, target_run)
    target_run.text = text
    for run in para.runs[1:]:
        run.text = ""


def _clear_run_color(run) -> None:
    """Remove explicit font color so Word uses the style default."""
    r_pr = run._element.get_or_add_rPr()
    for color_elem in r_pr.findall(qn("w:color")):
        r_pr.remove(color_elem)


def _copy_run_format(source, target) -> None:
    """Copy visible run formatting from source to target."""
    target.bold = source.bold
    target.italic = source.italic
    target.underline = source.underline
    if source.font.name:
        target.font.name = source.font.name
    if source.font.size:
        target.font.size = source.font.size
    if source.font.color and source.font.color.rgb:
        target.font.color.rgb = source.font.color.rgb


def _replace_in_cell(cell, replacements: dict[str, str]) -> None:
    for para in cell.paragraphs:
        _replace_in_paragraph(para, replacements)


def _append_cloned_row(table: Table, template_row: _Row) -> _Row:
    new_tr = deepcopy(template_row._tr)
    table._tbl.append(new_tr)
    return table.rows[-1]


def _iter_document_paragraphs(doc: docx.Document):
    """Yield every paragraph in the document body, tables, headers, and footers."""
    for para in doc.paragraphs:
        yield para
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    yield para
    for section in doc.sections:
        for para in section.header.paragraphs:
            yield para
        for para in section.footer.paragraphs:
            yield para


def _replace_in_paragraph(para, replacements: dict[str, str], count: int = 0) -> None:
    """Replace substrings in a paragraph while preserving run formatting where possible."""
    for old, new in replacements.items():
        if old not in para.text:
            continue
        _replace_once_multirun(para, old, new)
        if count == 1:
            break


def _replace_once_multirun(para, old: str, new: str) -> bool:
    """Replace one occurrence across runs without collapsing paragraph formatting."""
    runs = para.runs
    if not runs:
        if old in para.text:
            para.add_run(para.text.replace(old, new, 1))
            return True
        return False

    for run in runs:
        if old in run.text:
            run.text = run.text.replace(old, new, 1)
            return True

    full = para.text
    start = full.find(old)
    if start == -1:
        return False
    end = start + len(old)

    offsets: list[tuple[int, int]] = []
    for run_index, run in enumerate(runs):
        for char_index in range(len(run.text)):
            offsets.append((run_index, char_index))

    if not offsets or end > len(offsets):
        return False

    start_run, start_char = offsets[start]
    end_run, end_char = offsets[end - 1]

    prefix = runs[start_run].text[:start_char]
    suffix = runs[end_run].text[end_char + 1 :]

    runs[start_run].text = prefix + new
    for run_index in range(start_run + 1, end_run):
        runs[run_index].text = ""
    if end_run != start_run:
        runs[end_run].text = suffix
    elif suffix:
        runs[start_run].text += suffix

    return True
